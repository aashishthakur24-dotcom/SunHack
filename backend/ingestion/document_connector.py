"""Document connector — extracts text from PDF, DOCX, and plain files."""
from __future__ import annotations
import os


def extract_text(file_path: str) -> tuple[str, dict]:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file_path)
    else:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read(), {"title": os.path.basename(file_path)}


def _extract_pdf(path: str) -> tuple[str, dict]:
    import pdfplumber
    texts = []
    meta = {}
    with pdfplumber.open(path) as pdf:
        meta["pages"] = len(pdf.pages)
        meta["title"] = pdf.metadata.get("Title", os.path.basename(path))
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                texts.append(t)
    return "\n\n".join(texts), meta


def _extract_docx(path: str) -> tuple[str, dict]:
    from docx import Document
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    meta = {"title": os.path.basename(path), "paragraphs": len(doc.paragraphs)}
    return text, meta
